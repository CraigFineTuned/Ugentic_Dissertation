import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_style_font(style, font_name='Times New Roman', size=12, bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold

def create_custom_styles(doc):
    # Ensure Heading 1 is standard
    styles = doc.styles
    
    # Heading 1
    h1 = styles['Heading 1']
    set_style_font(h1, size=16, bold=True)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    
    # Heading 2
    h2 = styles['Heading 2']
    set_style_font(h2, size=14, bold=True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    
    # Heading 3
    h3 = styles['Heading 3']
    set_style_font(h3, size=12, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Normal Text
    normal = styles['Normal']
    set_style_font(normal, size=12)
    normal.paragraph_format.line_spacing = 1.5 # Academic standard

def generate_improved_docx(markdown_path, output_path):
    doc = Document()
    create_custom_styles(doc)

    # --- Front Matter ---
    doc.add_heading('LIST OF FIGURES', level=1)
    # (Ideally we'd parse the figures from the MD, but copying the placeholder block for now)
    figures = [
        'Figure 4.1: UGENTIC Multi-Agent System Architecture.........................................45',
        'Figure 4.2: Design Science Research Methodology Process.................................46',
        'Figure 5.1: Participant Distribution by Organizational Level (N=14)......................53',
        'Figure 5.2: Empirical Support for Major Themes Across Participant Sample (N=14).....55',
        'Figure 5.3: Research Question Coverage by Thematic Findings..........................65'
    ]
    for fig in figures:
        doc.add_paragraph(fig)
    doc.add_page_break()

    doc.add_heading('LIST OF ABBREVIATIONS', level=1)
    abbreviations = [
        'AI ................ Artificial Intelligence',
        'API ............... Application Programming Interface',
        'DSR ............. Design Science Research',
        'HR ................ Human Resources',
        'ICT ............... Information and Communication Technology',
        'IT .................. Information Technology',
        'ITIL ............... Information Technology Infrastructure Library',
        'ITSM ............ IT Service Management',
        'LLM ............. Large Language Model',
        'MAS ............. Multi-Agent System',
        'MCP ............. Model Context Protocol',
        'RAG ............. Retrieval-Augmented Generation',
        'ReAct ........... Reasoning and Acting',
        'RQ ................ Research Question',
        'RTA .............. Reflexive Thematic Analysis',
        'SAST ............ South African Standard Time',
        'SME ............. Small and Medium Enterprise',
        'UGENTIC ..... Ubuntu-Driven Agentic Collective Intelligence',
        'UK ................ United Kingdom',
        'UNESCO ...... United Nations Educational, Scientific and Cultural Organisation'
    ]
    for abbr in abbreviations:
        doc.add_paragraph(abbr)
    doc.add_page_break()
    
    doc.add_heading('KEY TERMS', level=1)
    # Key terms are paragraphs in the MD, handled in main loop or specifically here? 
    # The markdown has a "KEY TERMS" section. We'll let the main loop catch it or hardcode if needed.
    # The main loop logic below handles "KEY TERMS" as a header.

    # --- Main Parsing Logic ---
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split by newlines but keep blocks of text together
    lines = content.split('\n')
    
    # Regex for headings
    # Matches "1.1 BACKGROUND" or "1.1.1 The Disconnect"
    heading_regex = re.compile(r'^(\d+\.\d+(\.\d+)?)\s+(.*)')
    
    chapter_regex = re.compile(r'^CHAPTER\s+(\d+):\s+(.*)', re.IGNORECASE)
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line == '[PAGE BREAK]':
            doc.add_page_break()
            continue
        if line.startswith('**[PAGE') or line.startswith('---') or line.startswith('### **BLOCK'):
            continue

        # 1. Check for CHAPTER headings (Heading 1)
        chapter_match = chapter_regex.match(line)
        if chapter_match:
            # e.g., "CHAPTER 1: INTRODUCTION"
            doc.add_heading(line, level=1)
            continue

        # 2. Check for Special Headings (REFERENCES, APPENDICES)
        if line in ['REFERENCES', 'APPENDICES']:
            doc.add_heading(line, level=1)
            continue

        if line.startswith('Appendix A:'):
            doc.add_heading(line, level=2)
            continue

        # 3. Check for Key Terms / Lists (Heading 1 context from previous file)
        if line == 'KEY TERMS' or line == 'LIST OF FIGURES' or line == 'LIST OF ABBREVIATIONS':
            # Already handled explicitly above, or handled here if duplicative
            pass # We manually added them, so skip if encountered in text flow to avoid dupes?
                 # Actually, let's let the loop run. The MD has them.
                 # The MD file starts with [PAGE ix] blocks.
                 # Let's simpler logic: If we see a line that LOOKS like a header, style it.
        
        # 4. Check for Numbered Headings (Heading 2 & 3)
        # "1.1 BACKGROUND" -> Heading 2
        # "1.1.1 The Disconnect" -> Heading 3
        match = heading_regex.match(line)
        if match:
            numbering = match.group(1)
            title = match.group(3)
            level = numbering.count('.') + 1 # 1.1 = 1 dot = level 2? No.
            # 1.1 has 1 dot. Level 2.
            # 1.1.1 has 2 dots. Level 3.
            
            # Word Heading 1 is usually "Chapter".
            # Word Heading 2 is "1.1".
            # Word Heading 3 is "1.1.1".
            
            heading_level = min(level + 1, 3) # Cap at 3? Or strictly logic.
            if numbering.count('.') == 1: heading_level = 2
            if numbering.count('.') == 2: heading_level = 3
            
            doc.add_heading(line, level=heading_level)
            continue

        # 5. Check for bolded lines that might be headings (e.g., "**2.2 AI-ORGANIZATIONAL GAPS...**")
        if line.startswith('**') and line.endswith('**') and len(line) < 100:
            clean_line = line.replace('**', '').strip()
            # Check if it has numbers inside
            inner_match = heading_regex.match(clean_line)
            if inner_match:
                 # Same logic as above
                numbering = inner_match.group(1)
                if numbering.count('.') == 1: 
                    doc.add_heading(clean_line, level=2)
                elif numbering.count('.') == 2:
                    doc.add_heading(clean_line, level=3)
                else:
                    doc.add_heading(clean_line, level=2) # Fallback
                continue
            else:
                # Just a bold line? Could be a sub-sub-header or emphasis.
                # Let's make it bold text for now to be safe, or Heading 4 if we wanted deep nesting.
                p = doc.add_paragraph()
                p.add_run(clean_line).bold = True
                continue

        # 6. Special Case: KEY TERMS definitions
        # "**Term:** Definition"
        if line.startswith('**') and ':**' in line:
            # It's a key term.
            parts = line.split(':**', 1)
            term = parts[0].replace('**', '').strip()
            definition = parts[1].strip() if len(parts) > 1 else ""
            
            p = doc.add_paragraph()
            p.add_run(term + ':').bold = True
            p.add_run(' ' + definition)
            continue

        # 7. Figures
        if line.startswith('[INSERT FIGURE') or line.startswith('Figure '):
             p = doc.add_paragraph(line)
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             continue

        # 8. Normal Text
        doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Improved DOCX saved to: {output_path}")

if __name__ == '__main__':
    input_md = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Correction Report.md"
    output_docx = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Final_Dissertation_Improved.docx"
    generate_improved_docx(input_md, output_docx)
