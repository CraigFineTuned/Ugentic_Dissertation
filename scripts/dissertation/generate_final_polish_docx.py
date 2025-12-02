import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_style_font(style, font_name='Times New Roman', size=12, bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold

def create_custom_styles(doc):
    styles = doc.styles
    
    # Heading 1 (Chapter Titles)
    h1 = styles['Heading 1']
    set_style_font(h1, size=16, bold=True)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5
    
    # Heading 2 (1.1)
    h2 = styles['Heading 2']
    set_style_font(h2, size=14, bold=True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5
    
    # Heading 3 (1.1.1)
    h3 = styles['Heading 3']
    set_style_font(h3, size=12, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.5

    # Normal Text
    normal = styles['Normal']
    set_style_font(normal, size=12)
    normal.paragraph_format.line_spacing = 1.5 # 1.5 Spacing Requirement
    normal.paragraph_format.space_after = Pt(6) # Ensure paragraph separation

def scrub_text(text):
    """
    Applies the 'YOLO Polish' rules:
    1. Nuke Em-dashes
    2. Enforce SA English
    """
    # 1. Em-dashes
    text = text.replace('—', ' - ')
    text = text.replace('–', ' - ') # En-dash too, just in case

    # 2. SA English Dictionary
    replacements = {
        'Organization': 'Organisation',
        'Organizational': 'Organisational',
        'Behavior': 'Behaviour',
        'Labor': 'Labour',
        'Center': 'Centre',
        'Program': 'Programme',
        'Analyze': 'Analyse',
        'Operationalize': 'Operationalise',
        'Realize': 'Realise',
        'Prioritize': 'Prioritise',
        'Optimize': 'Optimise'
    }
    
    # Exceptions (Keep "AI-Organizational Gap" if it's a specific term, but user said "AI-Organizational can stay")
    # We will apply replacements but perform a negative lookbehind/lookahead for exceptions if needed?
    # Actually, simple replace is safer for "Behaviour", "Labour".
    # For "Organization", we need to be careful about "World Health Organization".
    
    for us, sa in replacements.items():
        if us == 'Organization':
             if 'World Health Organization' in text or 'International Labour Organization' in text:
                 pass # Skip blind replace, rely on manual check? 
                 # No, let's replace generic ones.
                 # Basic logic: Replace 'Organization ' with 'Organisation '
                 text = text.replace('Organization ', 'Organisation ')
                 text = text.replace('Organization.', 'Organisation.')
                 text = text.replace('Organization,', 'Organisation,')
        elif us == 'Organizational':
             # User said "Terms such as AI-Organizational can stay"
             if 'AI-Organizational' in text:
                 pass # Don't replace this specific instance
                 # We replace occurrences that are NOT preceded by 'AI-'
                 # Python replace doesn't do regex easily without import re, so let's use re
                 text = re.sub(r'(?<!AI-)Organizational', 'Organisational', text)
        elif us == 'Program':
            if 'Computer program' not in text:
                 text = text.replace(us, sa)
        else:
            text = text.replace(us, sa)
            text = text.replace(us.lower(), sa.lower())

    return text

def generate_final_polish_docx(markdown_path, output_path):
    doc = Document()
    create_custom_styles(doc)

    # --- Front Matter with Proper Spacing ---
    doc.add_heading('LIST OF FIGURES', level=1)
    figures = [
        'Figure 4.1: UGENTIC Multi-Agent System Architecture.........................................45',
        'Figure 4.2: Design Science Research Methodology Process.................................46',
        'Figure 5.1: Participant Distribution by Organizational Level (N=14)......................53',
        'Figure 5.2: Empirical Support for Major Themes Across Participant Sample (N=14).....55',
        'Figure 5.3: Research Question Coverage by Thematic Findings..........................65'
    ]
    for fig in figures:
        doc.add_paragraph(fig)
    doc.add_page_break()

    doc.add_heading('LIST OF ABBREVIATIONS', level=1)
    abbreviations = [
        'AI ................ Artificial Intelligence',
        'API ............... Application Programming Interface',
        'DSR ............. Design Science Research',
        'HR ................ Human Resources',
        'ICT ............... Information and Communication Technology',
        'IT .................. Information Technology',
        'ITIL ............... Information Technology Infrastructure Library',
        'ITSM ............ IT Service Management',
        'LLM ............. Large Language Model',
        'MAS ............. Multi-Agent System',
        'MCP ............. Model Context Protocol',
        'RAG ............. Retrieval-Augmented Generation',
        'ReAct ........... Reasoning and Acting',
        'RQ ................ Research Question',
        'RTA .............. Reflexive Thematic Analysis',
        'SAST ............ South African Standard Time',
        'SME ............. Small and Medium Enterprise',
        'UGENTIC ..... Ubuntu-Driven Agentic Collective Intelligence',
        'UK ................ United Kingdom',
        'UNESCO ...... United Nations Educational, Scientific and Cultural Organisation'
    ]
    for abbr in abbreviations:
        doc.add_paragraph(abbr)
    doc.add_page_break()
    
    # KEY TERMS (Definitions at beginning only)
    doc.add_heading('KEY TERMS', level=1)
    # (The MD file content loop will fill this in, we just set the header)

    # --- Main Parsing Logic ---
    with open(markdown_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    heading_regex = re.compile(r'^(\d+\.\d+(\.\d+)?)\s+(.*)')
    chapter_regex = re.compile(r'^CHAPTER\s+(\d+):\s+(.*)', re.IGNORECASE)
    
    for line in lines:
        line = line.strip()
        if not line: continue
        if line == '[PAGE BREAK]':
            doc.add_page_break()
            continue
        if line.startswith('**[PAGE') or line.startswith('---') or line.startswith('### **BLOCK'):
            continue
        if line in ['LIST OF FIGURES', 'LIST OF ABBREVIATIONS', 'KEY TERMS']: # handled manually
            continue

        # SCRUB TEXT BEFORE PROCESSING
        line = scrub_text(line)

        # 1. Chapters
        chapter_match = chapter_regex.match(line)
        if chapter_match:
            doc.add_heading(line.upper(), level=1) # Force Uppercase
            continue

        # 2. Special Headers
        if line in ['REFERENCES', 'APPENDICES']:
            doc.add_heading(line, level=1)
            continue
        if line.startswith('Appendix A:'):
            doc.add_heading(line, level=2)
            continue

        # 3. Numbered Headings
        match = heading_regex.match(line)
        if match:
            numbering = match.group(1)
            level = min(numbering.count('.') + 1, 3)
            doc.add_heading(line, level=level)
            continue

        # 4. Bold headings (sub-headers)
        if line.startswith('**') and line.endswith('**') and len(line) < 100 and ':**' not in line:
             clean = line.replace('**', '').strip()
             doc.add_heading(clean, level=2) # Treat as H2 or H3
             continue

        # 5. Key Terms (Definition format)
        if line.startswith('**') and ':**' in line:
            parts = line.split(':**', 1)
            term = parts[0].replace('**', '').strip()
            defn = parts[1].strip()
            p = doc.add_paragraph()
            runner = p.add_run(term + ':')
            runner.bold = True
            p.add_run(' ' + defn)
            continue

        # 6. Figures
        if line.startswith('[INSERT FIGURE') or line.startswith('Figure '):
             p = doc.add_paragraph(line)
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             continue

        # 7. Normal Paragraphs (No bullets allowed!)
        if line.startswith('* ') or line.startswith('- '):
            # Convert bullet to sentence? Or just remove bullet?
            # User said "No bullet points... try to use paragraphs."
            # If it's a list item, we append it to a paragraph or make it one.
            clean = line[2:].strip()
            doc.add_paragraph(clean) # Just add as text, lost the bullet.
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Final Polish DOCX saved to: {output_path}")

if __name__ == '__main__':
    input_md = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Correction Report.md"
    output_docx = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Final_Dissertation_YOLO.docx"
    generate_final_polish_docx(input_md, output_docx)
