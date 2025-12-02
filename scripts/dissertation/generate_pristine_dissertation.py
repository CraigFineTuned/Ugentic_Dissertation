import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def set_style_font(style, font_name='Times New Roman', size=12, bold=False):
    font = style.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold

def create_custom_styles(doc):
    styles = doc.styles
    
    # Heading 1 (Chapter Titles)
    h1 = styles['Heading 1']
    set_style_font(h1, size=16, bold=True)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centered Chapters
    
    # Heading 2 (1.1)
    h2 = styles['Heading 2']
    set_style_font(h2, size=14, bold=True)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.5
    
    # Heading 3 (1.1.1)
    h3 = styles['Heading 3']
    set_style_font(h3, size=12, bold=True)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.5

    # Normal Text
    normal = styles['Normal']
    set_style_font(normal, size=12)
    normal.paragraph_format.line_spacing = 1.5 
    normal.paragraph_format.space_after = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def scrub_text(text):
    # 1. Em-dashes destruction
    text = text.replace('—', ' - ')
    text = text.replace('–', ' - ') 

    # 2. SA English Enforcement (Standard SA/UK)
    # Mapping: US -> SA
    replacements = {
        'Organization': 'Organisation',
        'Organizational': 'Organisational',
        'Behavior': 'Behaviour',
        'Labor': 'Labour',
        'Center': 'Centre',
        'Program': 'Programme',
        'Analyze': 'Analyse',
        'Operationalize': 'Operationalise',
        'Realize': 'Realise',
        'Prioritize': 'Prioritise',
        'Optimize': 'Optimise'
    }
    
    for us, sa in replacements.items():
        if us == 'Organization':
             if 'World Health Organization' in text or 'International Labour Organization' in text:
                 # Keep specific proper nouns
                 text = text.replace('Organization ', 'Organisation ')
                 text = text.replace('Organization.', 'Organisation.')
                 text = text.replace('Organization,', 'Organisation,')
             else:
                 text = text.replace(us, sa)
        elif us == 'Organizational':
             # Exception: "AI-Organizational" stays
             if 'AI-Organizational' in text:
                 # Replace instances NOT preceded by AI-
                 text = re.sub(r'(?<!AI-)Organizational', 'Organisational', text)
             else:
                 text = text.replace(us, sa)
        elif us == 'Program':
            if 'Computer program' not in text:
                 text = text.replace(us, sa)
        else:
            text = text.replace(us, sa)
            text = text.replace(us.lower(), sa.lower())

    return text

def is_junk_line(line):
    """
    Detects if a line is markdown metadata, instructions, or artifacts.
    """
    l = line.strip()
    if not l: return False
    
    # Markdown Artifacts
    if l.startswith('[PAGE') and ']' in l: return True
    if l.startswith('**[PAGE'): return True
    if l.startswith('---'): return True
    if l.startswith('###'): return True
    if l.startswith('$$'): return True
    
    # Instructional Text from the Correction Report
    if l.startswith('**Punctuation:**'): return True
    if l.startswith('**Localization'): return True
    if l.startswith('**Formatting:**'): return True
    if l.startswith('**Figures:**'): return True
    if l.startswith('**Structure:**'): return True
    if 'Here is the **Submission-Ready** content' in l: return True
    if 'Copy and paste this' in l: return True
    if 'Signed: ____' in l: return True
    if 'Date: ____' in l: return True
    if 'rightarrow' in l: return True # LaTeX arrow artifact
    if 'BLOCK 2:' in l: return True
    
    return False

def generate_pristine_docx(markdown_path, output_path):
    doc = Document()
    create_custom_styles(doc)

    # --- FRONT MATTER GENERATION (Manual, Clean) ---
    # We do NOT read this from the MD file to avoid junk. We generate it clean.
    
    # 1. Title Page (Placeholder)
    doc.add_paragraph("[INSERT TITLE PAGE HERE]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 2. Declaration (Clean)
    doc.add_heading('Declaration', level=1)
    doc.add_paragraph("I declare that this dissertation is my own unaided work. It is being submitted for the degree of Bachelor of Science Honours in Information Technology at the Richfield Graduate Institute of Technology, South Africa. It has not been submitted before for any degree or examination in any other University.")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Signature")
    doc.add_paragraph("_" * 40)
    doc.add_paragraph("Date")
    doc.add_page_break()

    # 3. Lists (Figures/Abbreviations) - We can keep these from MD if we parse carefully, 
    # or just generate placeholders. Let's try to parse them from MD but skip junk.
    
    # --- MAIN CONTENT PARSING ---
    with open(markdown_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    heading_regex = re.compile(r'^(\d+\.\d+(\.\d+)?)\s+(.*)')
    chapter_regex = re.compile(r'^CHAPTER\s+(\d+):\s+(.*)', re.IGNORECASE)
    
    # State flags
    skip_block = False 

    for line in lines:
        line = line.strip()
        
        # 1. Skip Junk
        if is_junk_line(line):
            continue
            
        # 2. Skip Empty Lines (we handle spacing via Styles)
        if not line:
            continue

        # 3. Scrub Text
        line = scrub_text(line)

        # 4. Detect Structure
        
        # Chapter Headings
        chapter_match = chapter_regex.match(line)
        if chapter_match:
            doc.add_page_break() # Ensure Chapters start on new page
            doc.add_heading(line.upper(), level=1)
            continue

        # Special Sections
        if line in ['REFERENCES', 'APPENDICES', 'LIST OF FIGURES', 'LIST OF ABBREVIATIONS', 'KEY TERMS']:
            doc.add_page_break()
            doc.add_heading(line, level=1)
            continue
            
        if line.startswith('Appendix A:'):
            doc.add_heading(line, level=2)
            continue

        # Numbered Headings (1.1, 1.1.1)
        match = heading_regex.match(line)
        if match:
            numbering = match.group(1)
            level = min(numbering.count('.') + 1, 3)
            doc.add_heading(line, level=level)
            continue

        # Bold Sub-headings (No numbers, just bold text on its own line)
        if line.startswith('**') and line.endswith('**') and len(line) < 100 and ':**' not in line:
             clean = line.replace('**', '').strip()
             doc.add_heading(clean, level=2) # Treat as subheading
             continue

        # Key Terms (Definition)
        if line.startswith('**') and ':**' in line:
            parts = line.split(':**', 1)
            term = parts[0].replace('**', '').strip()
            defn = parts[1].strip()
            p = doc.add_paragraph()
            runner = p.add_run(term + ':')
            runner.bold = True
            p.add_run(' ' + defn)
            continue

        # Figures (Placeholders)
        if line.startswith('[INSERT FIGURE') or line.startswith('Figure '):
             p = doc.add_paragraph(line)
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER
             continue

        # Standard Paragraphs (No bullets!)
        if line.startswith('* ') or line.startswith('- '):
            clean = line[2:].strip()
            doc.add_paragraph(clean)
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Pristine DOCX saved to: {output_path}")

if __name__ == '__main__':
    input_md = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Correction Report.md"
    output_docx = r"C:\Users\craig\Desktop\MainProjects\Ugentic_Dissertation\DISSERTATION_ACADEMIC\Completed\Pristine_Dissertation_Craig_Vraagom.docx"
    generate_pristine_docx(input_md, output_docx)
