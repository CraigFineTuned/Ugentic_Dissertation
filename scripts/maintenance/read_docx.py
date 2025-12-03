import sys
import docx
import os

# Set stdout to handle utf-8
sys.stdout.reconfigure(encoding='utf-8')

def read_docx(file_path):
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        return

    try:
        doc = docx.Document(file_path)
        full_text = []
        
        print(f"--- START OF DOCUMENT: {os.path.basename(file_path)} ---")
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract tables (basic extraction)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))

        print("\n".join(full_text))
        print(f"\n--- END OF DOCUMENT ---")

    except Exception as e:
        print(f"Error reading .docx file: {e}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_docx.py <path_to_docx>")
    else:
        read_docx(sys.argv[1])