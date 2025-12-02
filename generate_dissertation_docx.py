from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_dissertation_docx(markdown_file_path, output_docx_path):
    document = Document()

    # --- Preliminary Sections ---
    # Add List of Figures placeholder
    document.add_heading('LIST OF FIGURES', level=1)
    document.add_paragraph('Figure 4.1: UGENTIC Multi-Agent System Architecture.........................................45')
    document.add_paragraph('Figure 4.2: Design Science Research Methodology Process.................................46')
    document.add_paragraph('Figure 5.1: Participant Distribution by Organizational Level (N=14)......................53')
    document.add_paragraph('Figure 5.2: Empirical Support for Major Themes Across Participant Sample (N=14).....55')
    document.add_paragraph('Figure 5.3: Research Question Coverage by Thematic Findings..........................65')
    document.add_page_break()

    # Add List of Abbreviations placeholder
    document.add_heading('LIST OF ABBREVIATIONS', level=1)
    document.add_paragraph('AI ................ Artificial Intelligence')
    document.add_paragraph('API ............... Application Programming Interface')
    document.add_paragraph('DSR ............. Design Science Research')
    document.add_paragraph('HR ................ Human Resources')
    document.add_paragraph('ICT ............... Information and Communication Technology')
    document.add_paragraph('IT .................. Information Technology')
    document.add_paragraph('ITIL ............... Information Technology Infrastructure Library')
    document.add_paragraph('ITSM ............ IT Service Management')
    document.add_paragraph('LLM ............. Large Language Model')
    document.add_paragraph('MAS ............. Multi-Agent System')
    document.add_paragraph('MCP ............. Model Context Protocol')
    document.add_paragraph('RAG ............. Retrieval-Augmented Generation')
    document.add_paragraph('ReAct ........... Reasoning and Acting')
    document.add_paragraph('RQ ................ Research Question')
    document.add_paragraph('RTA .............. Reflexive Thematic Analysis')
    document.add_paragraph('SAST ............ South African Standard Time')
    document.add_paragraph('SME ............. Small and Medium Enterprise')
    document.add_paragraph('UGENTIC ..... Ubuntu-Driven Agentic Collective Intelligence')
    document.add_paragraph('UK ................ United Kingdom')
    document.add_paragraph('UNESCO ...... United Nations Educational, Scientific and Cultural Organisation')
    document.add_page_break()

    # Add Key Terms placeholder
    document.add_heading('KEY TERMS', level=1)
    document.add_paragraph('AI-Organizational Gap: The misalignment between AI system capabilities, such as individual optimization and autonomous decision-making, and organisational operational realities, such as collective coordination and hierarchical consultation, which prevents effective AI integration despite technical sophistication.')
    document.add_paragraph('Bridging Mechanism: A framework connecting AI technical capabilities with organisational collaboration needs while respecting both computational constraints and cultural values. Ubuntu philosophy is investigated as a potential bridging mechanism in this research.')
    document.add_paragraph('Ubuntu Philosophy: A Southern African philosophical framework emphasizing collective humanity, mutual responsibility, and relational existence. It is encapsulated in the phrase umuntu ngumuntu ngabantu ("a person is a person through other people") and is used in this research as a potential cultural bridge aligning AI behaviours with organisational collaboration needs.')
    document.add_paragraph('UGENTIC: Ubuntu-Driven Agentic Collective Intelligence. This is a research instrument, not a commercial product, comprising six AI agents mirroring the GrandWest IT department structure. It implements Ubuntu principles as potential bridging mechanisms to enable the empirical investigation of bridging effectiveness through stakeholder assessment.')
    document.add_paragraph('Design Science Research (DSR): A research methodology investigating problems through artifact design and evaluation. UGENTIC serves as the designed artifact enabling the systematic investigation of Ubuntu as a bridging mechanism, with evaluation conducted through stakeholder assessment rather than technical performance metrics.')
    document.add_paragraph('Reflexive Thematic Analysis: A qualitative analysis methodology that emphasizes researcher reflexivity and participant meaning-making. It is used to analyse stakeholder experiences of Ubuntu-AI bridging effectiveness, respecting diverse interpretations across organisational levels.')
    document.add_page_break()

    # --- Main Content Parsing ---
    with open(markdown_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_chapter_num = 0
    in_references_section = False

    for line in lines:
        stripped_line = line.strip()

        if stripped_line == '[PAGE BREAK]':
            document.add_page_break()
            continue
        
        # Skip special instruction blocks in the markdown source
        if stripped_line.startswith('**[PAGE') or stripped_line.startswith('---'):
            continue
        if stripped_line.startswith('### **BLOCK'): # Skip the block indicators
            continue
        
        # Main Chapter Headings (e.g., CHAPTER 1: INTRODUCTION)
        if stripped_line.startswith('CHAPTER '):
            chapter_title = stripped_line.replace('CHAPTER ', '', 1).strip()
            # Extract chapter number for section numbering
            try:
                current_chapter_num = int(chapter_title.split(':')[0].strip())
            except ValueError:
                current_chapter_num = 0 # Reset or handle non-numeric chapters
            document.add_heading(chapter_title, level=1)
            in_references_section = (chapter_title == 'REFERENCES')
        elif stripped_line.startswith('**'): # Sub-subheadings and bolded text
            text_content = stripped_line.replace('**', '').strip()
            if text_content.startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
                # This could be a chapter section like '1.1 BACKGROUND'
                # Need to differentiate from simple bolded text
                # For now, treat all bolded sections as potential headings and refine
                document.add_heading(text_content, level=2)
            elif stripped_line.startswith('***'): # More specific case for sub-subheadings if needed. 
                document.add_heading(text_content, level=3)
            else: # Just bolded text within a paragraph
                p = document.add_paragraph()
                p.add_run(text_content).bold = True
        elif stripped_line.startswith(tuple(f'{i}.' for i in range(1, 7))): # Numerical headings like 1.1, 1.1.1
            parts = stripped_line.split(' ', 1)
            if len(parts) > 1:
                prefix = parts[0]
                text = parts[1]
                if prefix.count('.') == 1: # e.g., 1.1, 2.2
                    document.add_heading(stripped_line, level=2)
                elif prefix.count('.') == 2: # e.g., 1.1.1, 2.2.1
                    document.add_heading(stripped_line, level=3)
                else: # Default paragraph for other numbered lists if not a heading
                    document.add_paragraph(stripped_line)
            else:
                document.add_paragraph(stripped_line)
        elif stripped_line: # Regular paragraph
            if stripped_line.startswith('[INSERT FIGURE'):
                document.add_paragraph(stripped_line, style='Normal') # Keep placeholder as is for user
            elif in_references_section:
                # References often have hanging indents and specific formatting.
                # For now, treat as normal paragraph, user can apply hanging indent manually.
                document.add_paragraph(stripped_line, style='Normal')
            else:
                document.add_paragraph(stripped_line)
        else: # Empty line, adds spacing
            document.add_paragraph('')

    # --- Appendices ---
    document.add_page_break()
    document.add_heading('APPENDICES', level=1)
    document.add_heading('Appendix A: Ethical Clearance Letter', level=2)
    document.add_paragraph('[INSERT YOUR SCANNED ETHICAL CLEARANCE LETTER HERE]')


    document.save(output_docx_path)
    print(f"Generated DOCX saved to: {output_docx_path}")

if __name__ == "__main__":
    markdown_source = "C:\\Users\\craig\\Desktop\\MainProjects\\Ugentic_Dissertation\\DISSERTATION_ACADEMIC\\Completed\\Correction Report.md"
    output_target = "C:\\Users\\craig\\Desktop\\MainProjects\\Ugentic_Dissertation\\DISSERTATION_ACADEMIC\\Completed\\Final_Dissertation_From_Gemini.docx"
    
    generate_dissertation_docx(markdown_source, output_target)
    print("\nGeneration complete. Please review the 'Final_Dissertation_From_Gemini.docx' for formatting and placeholders.")
    print("Manual steps remaining: Insert actual images, update Table of Contents, sign declaration, insert ethical clearance.")
